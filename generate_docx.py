import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_business_plan_docx(filename):
    doc = Document()
    
    # Page Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles
    navy = RGBColor(15, 23, 42)
    amber = RGBColor(217, 119, 6)
    dark_gray = RGBColor(51, 65, 85)
    
    # Document Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("FORRETNINGSPLAN")
    r_title.bold = True
    r_title.font.size = Pt(26)
    r_title.font.color.rgb = navy
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("SaaS Coupon Spy — The «Honey» for SaaS & AI Subscriptions")
    r_sub.bold = True
    r_sub.font.size = Pt(15)
    r_sub.font.color.rgb = amber
    p_sub.paragraph_format.space_after = Pt(14)
    
    # Metadata Box
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Dokumentversjon:", "1.0 (Offisiell lanseringsplan)"),
        ("Selskap:", "SaaS Coupon Spy / aiappsy"),
        ("Forfatter / Eier:", "Pål Alexander Juritzen"),
        ("Dato:", "August 2026")
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.text = k
        c2.text = v
        c1.paragraphs[0].runs[0].bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        c2.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c1, "F1F5F9")
        set_cell_background(c2, "F8FAFC")
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
        c1.width = Inches(2.0)
        c2.width = Inches(4.5)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    def add_h1(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = navy
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = amber
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10.5)
        r.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = navy
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

    # 1. Executive Summary
    add_h1("1. Sammendrag (Executive Summary)")
    add_body("SaaS Coupon Spy er en automatisert nettleserutvidelse (Chrome Extension) og skyplattform som fungerer som et «Honey» dedikert utelukkende til SaaS (Software-as-a-Service), AI-verktøy, webhotell og digitale abonnementer.")
    add_body("Når brukere er på betalingssiden til verktøy som Canva, Hostinger, ElevenLabs, HeyGen, Cursor, Adobe eller NordVPN, finner og tester utvidelsen automatisk verifiserte rabattkoder på 2 sekunder – og aktiverer samtidig høyt betalende affiliate-provisjoner for selskapet.")
    
    add_h2("1.1 Visjon & Hovedmålsetting")
    add_bullet(" Nå 25 000 aktive månedlige brukere (MAU) og $12 500 i månedlig inntekt (MRR).", "År 1:")
    add_bullet(" Skalere til 100 000 aktive brukere og etablere direkte sponsoravtaler med ledende SaaS-selskaper.", "År 2:")
    add_bullet(" Nå 350 000 aktive brukere med en årlig omsetning (ARR) på over $2,2 millioner og bruttomargin over 95 %.", "År 3:")

    # 2. Markedsanalyse
    add_h1("2. Markedsanalyse & Mulighetsrom")
    add_body("Det globale SaaS-markedet er verdsatt til over $310 milliarder, med over 1,2 milliarder årlige programvaretransaksjoner.")
    add_h2("2.1 Markedets Hovedproblemer:")
    add_bullet(" Honey og RetailMeNot fokuserer på klær, sko og fysiske varer. De mangler koder for moderne AI- og skaperverktøy.", "Tradisjonelle kupong-utvidelser svikter på SaaS:")
    add_bullet(" Skapere, utviklere og bedrifter bruker $150–$600 månedlig på abonnementer og søker aktivt etter rabattkoder.", "Høye løpende kostnader:")
    add_bullet(" 90 % av koder på generiske rabatt-nettsider er utløpte. SaaS Coupon Spy løser dette med automatisert AI-verifisering.", "Frustrerende ugyldige koder:")

    # 3. Produkt & Teknologi
    add_h1("3. Produkt & Teknologisk Konkurransefortrinn")
    add_bullet(" Injiserer koder direkte i Stripe, Paddle, Chargebee, LemonSqueezy og Shopify uten at siden må lastes inn på nytt.", "2-Sekunders DOM Auto-Tester:")
    add_bullet(" Elegant overlegg i isolert Shadow-DOM som aldri krasjer eller forstyrrer nettsiders styling.", "Shadow-DOM Floating Pill:")
    add_bullet(" Automatisk skanning og ekstrahering av nye rabattkoder fra e-poster, nyhetsbrev og Twitter/X ved hjelp av Gemini Flash AI.", "Gemini Flash AI Deals-Motor:")
    add_bullet(" Nye rabattkoder og affiliate-lenker distribueres globalt til alle brukere på millisekunder uten oppdatering i Chrome Web Store.", "Dynamisk Edge API:")

    # 4. Inntektsmodell
    add_h1("4. Inntektsmodell & Monetiseringsstrategi")
    
    t_rev = doc.add_table(rows=5, cols=3)
    t_rev.alignment = WD_TABLE_ALIGNMENT.CENTER
    rev_headers = ["Inntektskilde", "Beskrivelse", "Gjennomsnittlig Inntekt"]
    for j, h_text in enumerate(rev_headers):
        c = t_rev.rows[0].cells[j]
        c.text = h_text
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = navy
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c, "E2E8F0")
        set_cell_margins(c, top=80, bottom=80, left=120, right=120)
        
    rev_rows = [
        ("1. High-Ticket Affiliate CPA", "Engangsutbetaling når bruker kjøper hosting, VPN eller SEO-verktøy.", "$35 – $150 per salg"),
        ("2. Recurring Revshare MRR", "Månedlig provisjon på løpende abonnementer (Notion, ElevenLabs, Webflow).", "20 % – 50 % månedlig"),
        ("3. B2B Sponsorater", "Nye SaaS-startups betaler for 'Featured Deal'-plassering øverst i katalogen.", "$500 – $2 500 / mnd per tool"),
        ("4. SaaS Spy PRO (Valgfritt)", "Valgfritt premium-abonnement for brukere med tilgang til VIP-koder.", "$3.99 / mnd eller $29 / år")
    ]
    for i, (k1, k2, k3) in enumerate(rev_rows):
        row = t_rev.rows[i+1]
        for idx, val in enumerate([k1, k2, k3]):
            cell = row.cells[idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if idx == 0: cell.paragraphs[0].runs[0].bold = True
            if idx == 2: 
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = amber
            set_cell_background(cell, "FFFFFF" if i%2==0 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    # 5. Finansielle Projeksjoner (Tabell)
    add_h1("5. Finansielle Projeksjoner & 3-Års Budsjett")
    
    t_fin = doc.add_table(rows=8, cols=4)
    t_fin.alignment = WD_TABLE_ALIGNMENT.CENTER
    fin_headers = ["Nøkkeltall / Metrikk", "År 1", "År 2", "År 3"]
    for j, h_text in enumerate(fin_headers):
        c = t_fin.rows[0].cells[j]
        c.text = h_text
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = navy
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c, "E2E8F0")
        set_cell_margins(c, top=80, bottom=80, left=120, right=120)

    fin_rows = [
        ("Aktive Brukere (Avsluttende)", "25 000", "100 000", "350 000"),
        ("Gjennomførte Kjøp per År", "18 000", "95 000", "380 000"),
        ("Affiliate-inntekter (CPA & MRR)", "$86 400", "$494 000", "$2 090 000"),
        ("Sponsorater & PRO-abonnement", "$6 000", "$38 000", "$145 000"),
        ("Totale Bruttoinntekter", "$92 400", "$532 000", "$2 235 000"),
        ("Totale Driftskostnader (OPEX)", "($5 800)", "($28 500)", "($78 000)"),
        ("Netto Driftsresultat (EBITDA)", "$86 600", "$503 500", "$2 157 000")
    ]
    for i, row_data in enumerate(fin_rows):
        row = t_fin.rows[i+1]
        for idx, val in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if idx == 0: cell.paragraphs[0].runs[0].bold = True
            if idx > 0 and i in [4, 6]: 
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = amber if i==4 else RGBColor(16, 185, 129)
            set_cell_background(cell, "FFFFFF" if i%2==0 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    add_body("")
    add_h2("5.1 Lønnsomhet & Driftsmarginer:")
    add_body("Fordi Google Cloud Run og Gemini Flash API opererer med tilnærmet null marginalkostnad per brukerforespørsel, oppnår selskapet en netto driftsmargin (EBITDA) på over 93 % i År 1 og opp mot 96,5 % i År 3.")

    # 6. GTM Markedsplan
    add_h1("6. Markeds- & Vekststrategi (Go-To-Market)")
    add_bullet(" Optimalisering for søk som 'Canva promo code', 'Hostinger coupon' og 'AI tool discount'.", "1. Chrome Web Store Organisk SEO:")
    add_bullet(" 20-sekunders skjermopptak som viser hvordan man kutter regningen med $30 på 2 sekunder.", "2. Virale Kortvideoer (TikTok & Shorts):")
    add_bullet(" Målrettet lanseringskampanje med mål om Top 3 Product of the Day.", "3. Product Hunt & Hacker News:")
    add_bullet(" Sponsorater og omtaler i ledende nyhetsbrev som Superhuman, The Rundown AI og TLDR.", "4. B2B & AI-nyhetsbrev:")

    doc.save(filename)
    print(f"Created: {filename}")

def create_sales_listing_docx(filename):
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    navy = RGBColor(15, 23, 42)
    amber = RGBColor(217, 119, 6)
    emerald = RGBColor(16, 185, 129)
    dark_gray = RGBColor(51, 65, 85)

    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("SALGSPROSPEKT & ANNONSE")
    r_title.bold = True
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = navy
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Nøkkelferdig AI & SaaS Kupong-plattform [Turnkey Digital Forretning]")
    r_sub.bold = True
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = amber
    p_sub.paragraph_format.space_after = Pt(14)

    def add_h1(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = navy
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10.5)
        r.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = navy
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

    add_body("Er du på jakt etter en skalerbar digital forretningsmodell med over 90 % fortjenestemargin og automatiserte gjentakende inntekter?")
    add_body("Nå har du muligheten til å overta SaaS Coupon Spy – en ferdigutviklet, toppmoderne nettleserutvidelse (Chrome Extension) og skybasert Admin-portal. Plattformen fungerer som et «Honey» for det raskt voksende markedet innen AI-verktøy, webhotell og programvareabonnementer.")
    add_body("Når brukerne handler hos Canva, Hostinger, ElevenLabs, HeyGen, Cursor eller NordVPN, finner utvidelsen automatisk verifiserte rabattkoder på 2 sekunder – samtidig som den genererer høye affiliate-provisjoner ($35–$150 per salg + 20–50 % månedlig gjentakende inntekt) til eieren.")

    # Pricing Box Table
    add_h1("Finansielle Betingelser & Prissetting")
    
    t_price = doc.add_table(rows=3, cols=3)
    t_price.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_headers = ["Element", "Anbefalt Pris", "Hva Kjøper Får"]
    for j, h_text in enumerate(p_headers):
        c = t_price.rows[0].cells[j]
        c.text = h_text
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = navy
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c, "E2E8F0")
        set_cell_margins(c, top=80, bottom=80, left=120, right=120)

    p_rows = [
        ("1. Kjøpesum for Virksomheten (Engangsbeløp)", "kr 69 000,- NOK\n($6 900 USD)", "100 % full eiendomsrett til Chrome Extension, Next.js Admin Hub, Gemini AI-motor, all kildekode, domener og 3-års forretningsplan."),
        ("2. Månedlig Drifts- & Supportavtale (SLA)", "kr 2 990,- / mnd\n($290 USD/mnd)", "Google Cloud Run hosting, Gemini AI API-kvoter, teknisk vedlikehold/feilretting, månedlig oppdatering av rabattkoder og prioritert support.")
    ]
    for i, row_data in enumerate(p_rows):
        row = t_price.rows[i+1]
        for idx, val in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if idx == 0: cell.paragraphs[0].runs[0].bold = True
            if idx == 1: 
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = amber if i==0 else emerald
            set_cell_background(cell, "FFFFFF" if i%2==0 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    add_h1("Hva Følger Med i Kjøpet (100 % Full Eiendomsrett):")
    add_bullet(" Komplett React 18 + Tailwind extension med 2-sekunders auto-tester, Shadow-DOM overlay og popup-katalog.", "1. Chrome Extension v1.0.0:")
    add_bullet(" Next.js 14 kontrollpanel med sanntids analyse, kupong-CMS og Gemini Flash AI deals-parser.", "2. Full Skybasert Admin Hub:")
    add_bullet(" Komplett markedsføringsstrategi, Go-To-Market plan og 3-års finansielle budsjetter.", "3. Dokumentasjon & Forretningsplan:")
    add_bullet(" Multi-stage Dockerfile for Google Cloud Run og full kildekode på GitHub.", "4. Produksjonsklar Container & Kode:")

    add_h1("Hva Inkluderes i den Månedlige Driftsavtalen (kr 2 990/mnd):")
    add_bullet(" Full drift og overvåking på Google Cloud Run med 99,9 % oppetid.", "Hosting & Servere:")
    add_bullet(" Løpende API-forbruk til Gemini Flash for rabattkode-analyse inkludert.", "AI API-kvoter:")
    add_bullet(" Oppdatering av DOM-selektorer dersom checkout-sider endrer design.", "Kontinuerlig Vedlikehold:")
    add_bullet(" Månedlig kvalitetssikring og påfyll av nye verifiserte SaaS- og AI-tilbud.", "Kupong-oppdateringer:")
    add_bullet(" Direkte tilgang til teknisk bistand og rådgivning.", "Prioritert Support:")

    add_h1("Kontakt & Overdragelse:")
    add_bullet(" Signering av kjøpekontrakt og overdragelsesavtale.", "1. Avtale:")
    add_bullet(" Overføring av GitHub-kildekode, tilganger og domener.", "2. Overføring:")
    add_bullet(" 1-til-1 digital onboarding (1 time) for full innføring i systemet.", "3. Onboarding:")
    add_body("\nKontakt for demo og oversendelse av fullt prospekt:\nE-post: paljuritzen@gmail.com | GitHub: github.com/aiappsy/saas-coupon-spy", bold=True)

    doc.save(filename)
    print(f"Created: {filename}")

if __name__ == "__main__":
    out_dir = r"C:\Users\paul\.gemini\antigravity\scratch\saas-coupon-spy"
    bp_path = os.path.join(out_dir, "Forretningsplan_SaaS_Coupon_Spy.docx")
    sales_path = os.path.join(out_dir, "Salgsannonse_SaaS_Coupon_Spy.docx")
    
    create_business_plan_docx(bp_path)
    create_sales_listing_docx(sales_path)
